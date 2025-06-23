"""
Word Document Export Utility
Exports research results to .docx format with proper formatting and preserved links
Supports mixed HTML and Markdown content
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import re
from datetime import datetime
import html
from typing import List, Dict, Any, Optional

# Try to import markdown2docx for better markdown handling
try:
    from markdown2docx import Markdown2Docx
    MARKDOWN2DOCX_AVAILABLE = True
except ImportError:
    MARKDOWN2DOCX_AVAILABLE = False


def create_research_report_docx(query, results, is_raw_data=False, source_data=None):
    """
    Create a Word document from research results
    
    Args:
        query: The research query
        results: The research results (HTML formatted)
        is_raw_data: Whether this is raw data mode or AI synthesis mode
        source_data: Source data dictionary (required for raw mode)
    
    Returns:
        Document: The Word document object
    """
    doc = Document()
    
    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"Research Report: {query}")
    title_run.bold = True
    title_run.font.size = Pt(24)
    
    # Add timestamp
    timestamp = doc.add_paragraph()
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp_run = timestamp.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    timestamp_run.font.size = Pt(11)
    
    # Add mode indicator
    mode_para = doc.add_paragraph()
    if is_raw_data:
        mode_para.add_run("Mode: Raw Data Analysis (LLM not available)")
    else:
        mode_para.add_run("Mode: AI-Synthesized Analysis")
    
    doc.add_paragraph()  # Spacing
    
    # Try to use markdown2docx if available and content is primarily markdown
    if MARKDOWN2DOCX_AVAILABLE and is_primarily_markdown(results):
        parse_with_markdown2docx(doc, results)
    else:
        # Parse the mixed HTML/Markdown results and convert to Word document
        parse_mixed_content_to_docx(doc, results)
    
    # For raw data mode, add the source data
    if is_raw_data and source_data:
        doc.add_paragraph()  # Spacing
        doc.add_paragraph("Raw Source Data:", style='Heading 2')
        
        total_sources = sum(len(sources) for sources in source_data.values())
        doc.add_paragraph(f"Total Sources: {total_sources}")
        
        for source_type, sources in source_data.items():
            if sources:
                source_type_name = source_type.replace('_sources', '').replace('_', ' ').title()
                doc.add_paragraph(f"{source_type_name}:", style='Heading 3')
                
                for i, source in enumerate(sources):
                    doc.add_paragraph(f"Source {i+1}:", style='Heading 4')
                    doc.add_paragraph(f"Title: {source.get('title', 'N/A')}")
                    
                    url = source.get('url', '')
                    if url:
                        para = doc.add_paragraph()
                        para.add_run('URL: ')
                        add_hyperlink(para, url, url)
                    
                    content = source.get('content', 'No content available.')
                    doc.add_paragraph(f"Content: {content}")
                    doc.add_paragraph()  # Spacing
    
    return doc


def is_primarily_markdown(content):
    """
    Check if content is primarily markdown format
    """
    # Count markdown elements
    markdown_elements = len(re.findall(r'^#{1,6}\s+', content, re.MULTILINE))
    markdown_elements += len(re.findall(r'^[-*+]\s+', content, re.MULTILINE))
    markdown_elements += len(re.findall(r'^\|.*\|$', content, re.MULTILINE))
    markdown_elements += len(re.findall(r'\*\*.*?\*\*', content))
    markdown_elements += len(re.findall(r'\*.*?\*', content))
    
    # Count HTML elements
    html_elements = len(re.findall(r'<[^>]+>', content))
    
    return markdown_elements > html_elements


def parse_with_markdown2docx(doc, content):
    """
    Parse content using markdown2docx library
    """
    try:
        # Convert HTML to markdown first
        markdown_content = html_to_markdown(content)
        
        # Create markdown converter instance
        markdown_converter = Markdown2Docx()
        markdown_converter.convert(markdown_content)
        temp_doc = markdown_converter.doc
        
        # Copy content from temp_doc to our main doc
        for element in temp_doc.element.body:
            doc.element.body.append(element)
            
    except Exception as e:
        # Fallback to manual parsing if markdown2docx fails
        parse_mixed_content_to_docx(doc, content)


def html_to_markdown(html_content):
    """
    Convert HTML to Markdown format
    """
    # Simple HTML to Markdown conversion
    content = html_content
    
    # Convert headings
    content = re.sub(r'<h1[^>]*>(.*?)</h1>', r'# \1', content, flags=re.DOTALL)
    content = re.sub(r'<h2[^>]*>(.*?)</h2>', r'## \1', content, flags=re.DOTALL)
    content = re.sub(r'<h3[^>]*>(.*?)</h3>', r'### \1', content, flags=re.DOTALL)
    content = re.sub(r'<h4[^>]*>(.*?)</h4>', r'#### \1', content, flags=re.DOTALL)
    
    # Convert bold and italic
    content = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', content, flags=re.DOTALL)
    content = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', content, flags=re.DOTALL)
    
    # Convert lists
    content = re.sub(r'<ul[^>]*>(.*?)</ul>', r'\1', content, flags=re.DOTALL)
    content = re.sub(r'<ol[^>]*>(.*?)</ol>', r'\1', content, flags=re.DOTALL)
    content = re.sub(r'<li[^>]*>(.*?)</li>', r'* \1', content, flags=re.DOTALL)
    
    # Convert links
    content = re.sub(r'<a[^>]*href="([^"]*)"[^>]*>(.*?)</a>', r'[\2](\1)', content, flags=re.DOTALL)
    
    # Convert paragraphs
    content = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', content, flags=re.DOTALL)
    
    # Convert divs
    content = re.sub(r'<div[^>]*>(.*?)</div>', r'\1', content, flags=re.DOTALL)
    
    # Clean up extra whitespace
    content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
    
    return content.strip()


def parse_mixed_content_to_docx(doc, content):
    """
    Parse mixed HTML and Markdown content and convert to Word document elements
    """
    # First, normalize the content
    content = normalize_content(content)
    
    # Split content into blocks
    blocks = split_content_into_blocks(content)
    
    for block in blocks:
        if not block.strip():
            continue
            
        # Handle different content types
        if is_markdown_heading(block):
            add_markdown_heading(doc, block)
        elif is_html_heading(block):
            add_html_heading(doc, block)
        elif is_markdown_list(block):
            add_markdown_list(doc, block)
        elif is_html_list(block):
            add_html_list(doc, block)
        elif is_markdown_table(block):
            add_markdown_table(doc, block)
        elif is_html_table(block):
            add_html_table(doc, block)
        elif is_markdown_code_block(block):
            add_markdown_code_block(doc, block)
        elif is_html_code_block(block):
            add_html_code_block(doc, block)
        elif is_horizontal_rule(block):
            add_horizontal_rule(doc)
        elif is_link(block):
            add_link(doc, block)
        elif is_bold_text(block):
            add_bold_text(doc, block)
        elif is_italic_text(block):
            add_italic_text(doc, block)
        else:
            # Regular text content
            add_regular_text(doc, block)


def normalize_content(content):
    """
    Normalize mixed HTML and Markdown content
    """
    # Remove outer div wrappers
    content = re.sub(r'<div[^>]*>', '', content)
    content = re.sub(r'</div>', '', content)
    
    # Decode HTML entities
    content = html.unescape(content)
    
    # Normalize line endings
    content = content.replace('\r\n', '\n').replace('\r', '\n')
    
    return content


def split_content_into_blocks(content):
    """
    Split content into logical blocks for processing
    """
    # Split by double newlines to separate blocks
    blocks = re.split(r'\n\s*\n', content)
    return [block.strip() for block in blocks if block.strip()]


def is_markdown_heading(text):
    """Check if text is a Markdown heading"""
    return re.match(r'^#{1,6}\s+', text)


def is_html_heading(text):
    """Check if text is an HTML heading"""
    return re.match(r'^<h[1-6][^>]*>', text)


def is_markdown_list(text):
    """Check if text is a Markdown list"""
    lines = text.split('\n')
    return all(re.match(r'^[\s]*[-*+]\s+', line) for line in lines if line.strip())


def is_html_list(text):
    """Check if text is an HTML list"""
    return '<ul>' in text or '<ol>' in text or text.strip().startswith('<li>')


def is_markdown_table(text):
    """Check if text is a Markdown table"""
    lines = text.split('\n')
    if len(lines) < 3:
        return False
    # Check for table header and separator
    header_match = re.match(r'^\|.*\|$', lines[0])
    separator_match = re.match(r'^\|[\s]*:?-+:?[\s]*\|', lines[1])
    return bool(header_match and separator_match)


def is_html_table(text):
    """Check if text is an HTML table"""
    return '<table>' in text or text.strip().startswith('<tr>')


def is_markdown_code_block(text):
    """Check if text is a Markdown code block"""
    return text.startswith('```') or re.match(r'^    ', text, re.MULTILINE)


def is_html_code_block(text):
    """Check if text is an HTML code block"""
    return '<code>' in text or '<pre>' in text


def is_horizontal_rule(text):
    """Check if text is a horizontal rule"""
    return re.match(r'^[-*_]{3,}$', text.strip()) or '<hr' in text


def is_link(text):
    """Check if text contains a link"""
    return 'http' in text and ('URL:' in text or re.search(r'https?://[^\s]+', text))


def is_bold_text(text):
    """Check if text is bold (Markdown or HTML)"""
    return text.startswith('**') or text.startswith('<strong>')


def is_italic_text(text):
    """Check if text is italic (Markdown or HTML)"""
    return text.startswith('*') or text.startswith('<em>')


def add_markdown_heading(doc, text):
    """Add a Markdown heading to the document"""
    match = re.match(r'^(#{1,6})\s+(.+)$', text)
    if match:
        level = len(match.group(1))
        heading_text = match.group(2)
        
        para = doc.add_paragraph()
        run = para.add_run(heading_text)
        run.bold = True
        
        # Set font size based on heading level
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(16)
        elif level == 3:
            run.font.size = Pt(14)
        elif level == 4:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(11)


def add_html_heading(doc, text):
    """Add an HTML heading to the document"""
    match = re.search(r'<h([1-6])[^>]*>(.*?)</h\1>', text, re.DOTALL)
    if match:
        level = int(match.group(1))
        heading_text = match.group(2).strip()
        
        para = doc.add_paragraph()
        run = para.add_run(heading_text)
        run.bold = True
        
        # Set font size based on heading level
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(16)
        elif level == 3:
            run.font.size = Pt(14)
        elif level == 4:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(11)


def add_markdown_list(doc, text):
    """Add a Markdown list to the document"""
    lines = text.split('\n')
    for line in lines:
        if line.strip():
            # Extract list item text
            match = re.match(r'^[\s]*[-*+]\s+(.+)$', line)
            if match:
                item_text = match.group(1)
                para = doc.add_paragraph()
                para.style = 'List Bullet'
                para.add_run(item_text)


def add_html_list(doc, text):
    """Add an HTML list to the document"""
    # Extract list items
    items = re.findall(r'<li[^>]*>(.*?)</li>', text, re.DOTALL)
    for item in items:
        para = doc.add_paragraph()
        para.style = 'List Bullet'
        para.add_run(item.strip())


def add_markdown_table(doc, text):
    """Add a Markdown table to the document"""
    lines = text.split('\n')
    if len(lines) < 3:
        return
    
    # Parse table structure
    table_data = []
    for line in lines:
        if line.strip() and not line.strip().startswith('|'):
            continue
        if '|' in line:
            # Split by | and clean up
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(cells)
    
    if not table_data:
        return
    
    # Create table
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Table Grid'
    
    # Fill table data
    for i, row in enumerate(table_data):
        for j, cell in enumerate(row):
            if i < len(table.rows) and j < len(table.rows[i].cells):
                table.rows[i].cells[j].text = cell


def add_html_table(doc, text):
    """Add an HTML table to the document"""
    # Extract table rows
    rows = re.findall(r'<tr[^>]*>(.*?)</tr>', text, re.DOTALL)
    if not rows:
        return
    
    # Parse first row to determine column count
    first_row_cells = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', rows[0], re.DOTALL)
    if not first_row_cells:
        return
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=len(first_row_cells))
    table.style = 'Table Grid'
    
    # Fill table data
    for i, row in enumerate(rows):
        cells = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', row, re.DOTALL)
        for j, cell in enumerate(cells):
            if i < len(table.rows) and j < len(table.rows[i].cells):
                table.rows[i].cells[j].text = cell.strip()


def add_markdown_code_block(doc, text):
    """Add a Markdown code block to the document"""
    # Remove code block markers
    if text.startswith('```'):
        lines = text.split('\n')
        if len(lines) >= 3:
            # Remove first and last lines (``` markers)
            code_content = '\n'.join(lines[1:-1])
        else:
            code_content = text
    else:
        # Indented code block
        code_content = re.sub(r'^    ', '', text, flags=re.MULTILINE)
    
    para = doc.add_paragraph()
    run = para.add_run(code_content)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)


def add_html_code_block(doc, text):
    """Add an HTML code block to the document"""
    # Extract code content
    code_match = re.search(r'<code[^>]*>(.*?)</code>', text, re.DOTALL)
    if not code_match:
        code_match = re.search(r'<pre[^>]*>(.*?)</pre>', text, re.DOTALL)
    
    if code_match:
        code_content = code_match.group(1).strip()
        para = doc.add_paragraph()
        run = para.add_run(code_content)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)


def add_horizontal_rule(doc):
    """Add a horizontal rule to the document"""
    para = doc.add_paragraph()
    para.add_run('_' * 50)


def add_link(doc, text):
    """Add a link to the document"""
    url_match = re.search(r'https?://[^\s]+', text)
    if url_match:
        url = url_match.group(0)
        para = doc.add_paragraph()
        
        if text.startswith('URL:'):
            para.add_run('URL: ')
            add_hyperlink(para, url, url)
        else:
            add_hyperlink(para, url, url)
    else:
        para = doc.add_paragraph()
        para.add_run(text)


def add_bold_text(doc, text):
    """Add bold text to the document"""
    # Handle Markdown bold
    if text.startswith('**'):
        match = re.match(r'^\*\*(.+?)\*\*$', text)
        if match:
            para = doc.add_paragraph()
            run = para.add_run(match.group(1))
            run.bold = True
            return
    
    # Handle HTML bold
    match = re.search(r'<strong[^>]*>(.*?)</strong>', text, re.DOTALL)
    if match:
        para = doc.add_paragraph()
        run = para.add_run(match.group(1).strip())
        run.bold = True
        return
    
    # Fallback to regular text
    add_regular_text(doc, text)


def add_italic_text(doc, text):
    """Add italic text to the document"""
    # Handle Markdown italic
    if text.startswith('*') and not text.startswith('**'):
        match = re.match(r'^\*(.+?)\*$', text)
        if match:
            para = doc.add_paragraph()
            run = para.add_run(match.group(1))
            run.italic = True
            return
    
    # Handle HTML italic
    match = re.search(r'<em[^>]*>(.*?)</em>', text, re.DOTALL)
    if match:
        para = doc.add_paragraph()
        run = para.add_run(match.group(1).strip())
        run.italic = True
        return
    
    # Fallback to regular text
    add_regular_text(doc, text)


def add_regular_text(doc, text):
    """Add regular text to the document with inline markdown formatting"""
    # Clean up any remaining HTML tags
    clean_text = re.sub(r'<[^>]+>', '', text)
    
    if clean_text.strip():
        para = doc.add_paragraph()
        
        # Process inline markdown formatting
        current_pos = 0
        while current_pos < len(clean_text):
            # Look for bold text **text**
            bold_match = re.search(r'\*\*(.*?)\*\*', clean_text[current_pos:])
            # Look for italic text *text* (but not **text**)
            italic_match = re.search(r'(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)', clean_text[current_pos:])
            
            if bold_match and (not italic_match or bold_match.start() < italic_match.start()):
                # Add text before bold
                if bold_match.start() > 0:
                    para.add_run(clean_text[current_pos:current_pos + bold_match.start()])
                
                # Add bold text
                bold_run = para.add_run(bold_match.group(1))
                bold_run.bold = True
                
                current_pos += bold_match.end()
            elif italic_match:
                # Add text before italic
                if italic_match.start() > 0:
                    para.add_run(clean_text[current_pos:current_pos + italic_match.start()])
                
                # Add italic text
                italic_run = para.add_run(italic_match.group(1))
                italic_run.italic = True
                
                current_pos += italic_match.end()
            else:
                # No more formatting, add remaining text
                para.add_run(clean_text[current_pos:])
                break


def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph
    """
    # Create the w:hyperlink tag
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), 'rId1')  # This will be set by the document
    
    # Create the w:r element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Join all the xml elements together
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)
    
    # Set hyperlink styling
    r.font.color.rgb = None  # Blue color
    r.font.underline = True
    
    return hyperlink


def save_docx_to_bytes(doc):
    """
    Save document to bytes for Streamlit download
    """
    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_docx_filename(query):
    """
    Generate a filename for the Word document
    """
    # Clean the query for filename
    clean_query = re.sub(r'[^\w\s-]', '', query)
    clean_query = re.sub(r'[-\s]+', '-', clean_query)
    clean_query = clean_query[:50]  # Limit length
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    return f"Research_Report_{clean_query}_{timestamp}.docx" 